from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

SOURCE = Path(r"C:\Users\wangz\Desktop\初稿_ai完成版订.docx")
OUTPUT = Path(r"C:\Users\wangz\Desktop\肠道菌群——完成版\CMPB_compact_multiscale_transcriptomics_3500_words.docx")
FONT = "Times New Roman"


def set_font(run, size=11, bold=None, italic=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), FONT)


def add_para(doc, text="", *, style=None, size=11, bold=False, italic=False,
             align=None, before=0, after=5, keep=False):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.keep_with_next = keep
    if align is not None:
        p.alignment = align
    if text:
        set_font(p.add_run(text), size=size, bold=bold, italic=italic)
    return p


def add_labelled(doc, label, text):
    p = add_para(doc, after=4)
    set_font(p.add_run(label), size=10.5, bold=True)
    set_font(p.add_run(text), size=10.5)
    return p


def add_heading(doc, text, level=1):
    size = 13 if level == 1 else 11.5
    return add_para(doc, text, size=size, bold=True, before=10 if level == 1 else 7,
                    after=4, keep=True)


def add_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ("Data layer", "Material", "Role in the pipeline")
    for cell, text in zip(table.rows[0].cells, headers):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        set_font(cell.paragraphs[0].add_run(text), size=9.5, bold=True)
    rows = (
        ("Bulk discovery", "4 mucosal cohorts; 326 target/inflamed and 88 reference samples", "Cross-cohort discovery, WGCNA, glycolysis intersection and workflow comparison"),
        ("External test", "GSE47908 (n=54) and GSE13367 (n=27)", "Locked evaluation of the final eight-gene model"),
        ("Cell-resolved", "GSE214695: 30,068 cells from 6 HC and 6 UC donors; GSE189184: 2 HC and 5 UC sections", "Cellular localization and spatial directionality"),
        ("Knowledge/structure", "gutMGene v2.0 and ancillary PPARG docking outputs", "Candidate context and structural prioritization"),
    )
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            set_font(p.add_run(text), size=8.8)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_font(run, size=9 if row is not table.rows[0] else 9.5, bold=(row is table.rows[0]))
    return table


doc = Document(SOURCE)

# Preserve the source document's page setup and styles, while replacing its manuscript body.
body = doc._element.body
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)

for style_name in ("Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
    try:
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)
    except KeyError:
        pass

title = add_para(
    doc,
    "A multiscale transcriptomic computational pipeline identifies epithelial PPARG dysregulation and structurally prioritizes lariciresinol in ulcerative colitis",
    size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8, keep=True,
)
add_para(doc, "Keywords: ulcerative colitis; multiscale transcriptomics; PPARG; machine learning; lariciresinol", size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

add_heading(doc, "Abstract")
add_labelled(doc, "Background: ", "Ulcerative colitis (UC) transcriptomes are heterogeneous across cohorts and cellular compartments, which complicates reproducible target discovery. We developed a compact multiscale computational workflow to connect bulk discovery, external testing, cell-resolved validation and knowledge-guided structural prioritization.")
add_labelled(doc, "Methods: ", "Four discovery mucosal cohorts (326 UC or inflamed-UC and 88 reference samples) were integrated with two locked external cohorts, donor-aware single-cell RNA sequencing, spatial transcriptomics and gutMGene v2.0 curation. Differential expression, WGCNA, glycolysis-focused intersection and 117 leakage-controlled feature-selection/classification workflows defined stable candidates.")
add_labelled(doc, "Results: ", "The workflow yielded 24 candidates and an eight-gene signature. PPARG was downregulated in all discovery cohorts, ranked second in selection stability (84.4%), and contributed to an externally tested ridge model (AUC 0.957 and 0.865). In 30,068 single cells from 12 donors, PPARG was epithelial-enriched and epithelial PPARG regulon activity was lower in UC (mean difference -2.104; FDR 0.0026); spatial data were directionally concordant. gutMGene connected lariciresinol to an Enterococcus faecalis context. In ancillary docking outputs, lariciresinol had the most favorable Vina scores among seven candidates for 7AWC and 3FUR (-6.9 and -7.0 kcal/mol), but failed uniform admission under stricter multi-method geometric quality control.")
add_labelled(doc, "Conclusion: ", "The pipeline identifies a reproducible epithelial PPARG-associated UC signal and prioritizes, but does not validate, lariciresinol as a structural candidate. The docking evidence is preliminary; molecular-dynamics, target-engagement and functional experiments remain required.")

add_heading(doc, "1. Introduction")
add_para(doc, "Ulcerative colitis is characterized by epithelial barrier injury, immune activation and altered gut microbial ecology [1,2]. Public mucosal transcriptomes provide a scalable resource for target discovery, but signals may be distorted by cohort effects, disease activity and shifts in cell composition. A useful computational study must therefore distinguish cross-cohort reproducibility from cell-type localization and from database-derived biological hypotheses.")
add_para(doc, "PPARG is a plausible node at this interface. It is expressed in colonic epithelium, reduced in UC in earlier work, and linked to epithelial inflammatory responses [4,5]. Metabolic programmes, including glycolysis, may further contextualize PPARG dysregulation, but transcriptomic co-enrichment cannot establish regulatory direction. Likewise, curated microbe-metabolite-host associations can nominate candidates but do not demonstrate abundance, exposure or binding in the analysed patients [7,8].")
add_para(doc, "We therefore assembled a multiscale transcriptomic computational pipeline that (i) identifies cross-cohort candidates from bulk mucosa, (ii) compares learning workflows without validation leakage, (iii) resolves the PPARG signal in single-cell and spatial data, and (iv) uses gutMGene with ancillary structural evidence to prioritize a microbe-linked metabolite. This framing emphasizes the pipeline and its evidence boundaries rather than a causal UC mechanism.")
add_para(doc, "The central claim is therefore deliberately bounded: integration of independent transcriptomic layers can identify a stable PPARG-associated epithelial signal and rank a linked metabolite for follow-up. It does not establish that a microbial metabolite regulates PPARG in patients, that PPARG mediates glycolysis, or that a candidate ligand is therapeutically active.")

add_heading(doc, "2. Materials and methods")
add_heading(doc, "2.1 Study design and data layers", level=2)
add_para(doc, "We reanalysed de-identified GEO datasets [9]. Discovery comprised GSE73661, GSE75214, GSE87466 and GSE107499; the latter contrasts inflamed with non-inflamed UC biopsies rather than UC with healthy controls [10-13]. GSE47908 and GSE13367 were retained for locked external evaluation [14,15]. GSE214695 was restricted to six healthy-control and six UC donors, and GSE189184 supplied two healthy-control and five UC Visium sections [16,17]. Donors, bulk samples and sections, rather than cells or spots, were the inferential units.")
add_table(doc)
add_heading(doc, "2.2 Bulk discovery and model evaluation", level=2)
add_para(doc, "Shared genes were filtered, log2(x+1)-transformed and ComBat-corrected for discovery integration [18]. Limma identified differentially expressed genes (Benjamini-Hochberg adjusted P<0.05 and |log2 fold change|>1) [19]; GO/KEGG enrichment summarized their functions [20,21]. Candidate genes were the intersection of differentially expressed genes, trait-associated WGCNA genes [22] and a 911-gene glycolysis compendium from MSigDB, KEGG and Reactome [21,23,24].")
add_para(doc, "The 24-gene panel was evaluated using 117 selector-classifier workflows. In four leave-one-cohort-out folds, ranking, feature selection, tuning and fitting occurred only in the training data. Stable genes required selection frequency >=0.50 and a consistent direction across cohorts. An eight-gene ridge model fitted on all discovery data was applied once, without refitting, to the locked external cohorts.")
add_para(doc, "For cell-resolved analyses, broad cell-type proportions were calculated per donor and differential activity was assessed with donor-level pseudobulk or donor summaries. Spatial spots were quality filtered and deconvolved with a single-cell-derived reference; exact tests enumerated the 21 possible allocations of five UC and two healthy-control sections. Multiple comparisons were controlled with the Benjamini-Hochberg procedure. These choices were intended to avoid cell- and spot-level pseudo-replication.")
add_heading(doc, "2.3 Cell-resolved, knowledge and structural contextualization", level=2)
add_para(doc, "Single-cell compositions and PPARG expression were summarized per donor; Hallmark glycolysis was scored with UCell and PPARG activity with DoRothEA/decoupleR [23,25-27]. Spatial analyses used section-level summaries and treated spot-level correlations as descriptive. CIBERSORT and xCell outputs were analysed within cohorts and pooled by random effects [28,29]. gutMGene v2.0 links were used only for candidate context.")

add_heading(doc, "3. Results")
add_heading(doc, "3.1 Cross-cohort discovery and constrained candidate selection", level=2)
add_para(doc, "Across four discovery cohorts, 681 genes met the differential-expression threshold (437 higher and 244 lower in the target/inflamed group). Enrichment highlighted inflammatory, bacterial-response, cytokine and chemokine programmes. PPARG was reduced (log2 fold change -1.294; adjusted P=8.26x10^-42), and its effect was negative in all cohorts (Hedges' g -1.974 to -2.514). Intersection with WGCNA and glycolysis evidence produced 24 candidates, of which PPARG was the sole overlap with the 117-gene gutMGene-related set (Fig. 2).")
add_para(doc, "The enrichment pattern provides context rather than a PPARG-specific mechanism: Gene Ontology terms included humoral immune response, response to bacterial molecules and leukocyte migration, while KEGG terms included cytokine-cytokine receptor interaction, IL-17, TNF and chemokine signalling. The 24-gene panel was consequently treated as a constrained discovery set, not as evidence that every member directly regulates glycolysis or responds to microbial metabolites.")
add_heading(doc, "3.2 Leakage-controlled learning retained an externally tested PPARG-containing signature", level=2)
add_para(doc, "The all-feature ridge workflow had a mean leave-one-cohort-out AUC of 0.987 (minimum 0.965). Eight genes met the stability criterion: PDE6A, PPARG, ADH6, LCN2, VLDLR, SLC2A3, TRPM6 and KDELR3; PPARG ranked second (84.4%). In locked external testing, the eight-gene ridge model achieved AUCs of 0.957 in GSE47908 and 0.865 in GSE13367, exceeding the 24-gene model and PPARG alone. These values support a reproducible molecular signature, not a deployment-ready diagnostic model, because the discovery panel was defined before the cross-cohort comparison and the external cohorts were small (Fig. 3).")
add_para(doc, "External discrimination was accompanied by Brier scores of 0.074 and 0.144, respectively. The corresponding calibration slopes were 1.093 (95% CI 0.527-1.658) and 0.909 (0.233-1.585), but the uncertainty in GSE13367 (n=27) was substantial. Thus, external testing supports transportability across the two held-out datasets, whereas prospective performance across clinical settings remains unknown.")
add_heading(doc, "3.3 Single-cell and spatial analyses localized the PPARG signal", level=2)
add_para(doc, "The single-cell dataset contained 30,068 cells from 12 donors. Epithelial cells were reduced in UC by 37.47 percentage points (FDR=3.97x10^-4), whereas plasma cells increased by 24.74 percentage points (FDR=0.0094). PPARG expression was concentrated in epithelial cells, and epithelial PPARG regulon activity was lower in UC (mean difference -2.104; FDR=0.0026). PPARG and glycolysis were co-enriched across epithelial cell types, but donor-level correlations and case-control glycolysis differences were not significant (Fig. 4).")
add_para(doc, "In seven spatial sections, epithelial contribution was lower and PPARG regulon activity was directionally lower in UC across all prespecified summaries; exact section-level P values were 0.143. Within epithelial-enriched spots, PPARG activity and glycolysis were positively correlated (rho 0.123-0.536), but spatial autocorrelation precludes patient-level inference. Bulk deconvolution also showed broad immune remodelling, while PPARG was positively associated with epithelial and negatively associated with plasma-cell xCell scores across cohorts (pooled rho 0.497 and -0.473, respectively) (Figs. 5-6).")
add_para(doc, "The section-level UC-HC difference in PPARG activity was -0.455 for all spots, -0.616 in the epithelial-fraction upper quartile and -0.386 for the epithelial-weighted summary. CIBERSORT meta-analysis further indicated increased M0 macrophages, activated CD4 memory T cells and neutrophils, and reduced M2 macrophages, CD8 T cells and regulatory T cells in the target/inflamed group. The plasma-cell estimate was positive but imprecise. These composition-sensitive analyses are consistent with, but cannot by themselves explain, the epithelial PPARG signal.")
add_heading(doc, "4. Discussion")
add_para(doc, "This study reframes the contribution as a multiscale transcriptomic computational pipeline. Its decisive output is not a single causal mechanism but an evidence chain: consistent bulk PPARG suppression, leakage-aware model selection, epithelial localization, directionally concordant spatial evidence and a traceable microbe-metabolite context. This organization makes the strength of each inference explicit and is appropriate for computational biomedicine, where data integration can be valuable without substituting for experimental validation.")
add_para(doc, "The results extend prior reports of reduced epithelial PPARG in UC [5]. The single-cell regulon result is stronger than epithelial PPARG expression alone, but it remains a predicted transcriptional footprint, not direct binding or causal regulation. Similarly, epithelial PPARG-glycolysis co-enrichment is compatible with a metabolically altered epithelial state, yet the absence of significant donor-level correlation or broad case-control glycolysis differences rules out a claim that PPARG controls glycolysis.")
add_para(doc, "The model-comparison layer also illustrates a general pipeline lesson. High leave-one-cohort-out AUCs can be useful for selecting among workflows when preprocessing and tuning remain inside the training folds, but they do not convert discovery-derived panels into independent biomarkers. The locked external tests provide a separate check, although their small sample sizes limit calibration claims. Retaining this distinction prevents an apparently high-performing transcriptomic classifier from being overstated as a clinical diagnostic tool.")
add_para(doc, "The docking figures support a deliberately modest structural statement. Lariciresinol ranked best by Vina against two PPARG structures, while the common geometric quality gate did not admit it consistently across alternative docking routes. The appropriate interpretation is therefore structural prioritization with an explicit quality-control warning, not validated binding, agonism or therapeutic potential. No molecular-dynamics trajectory, convergence metric or free-energy result was supplied; molecular dynamics is consequently not represented as completed positive evidence in this manuscript.")
add_para(doc, "Several limitations remain. The analysis is retrospective and susceptible to residual platform, treatment and compositional effects; GSE107499 contributes an inflammation-within-UC contrast. Single-cell and spatial sample sizes limit donor- and section-level power, and deconvolution, gene-set and regulon scores are model-based. The cohorts lack matched microbiome, metabolite, protein and target-engagement measurements. Prospective multimodal cohorts, biochemical binding assays, epithelial PPARG perturbation and replicated molecular dynamics would be needed to test the proposed host-microbe-metabolite axis.")

add_heading(doc, "5. Conclusion")
add_para(doc, "A multiscale transcriptomic computational pipeline identified a reproducibly suppressed epithelial PPARG-associated signal in UC and an externally tested eight-gene signature. gutMGene and ancillary docking outputs prioritize lariciresinol, but strict multi-method quality control prevents a binding claim. The manuscript therefore provides a compact, evidence-bounded computational framework for candidate prioritization rather than a completed mechanistic or therapeutic validation.")

add_heading(doc, "Data and code availability")
add_para(doc, "All analysed expression datasets are publicly available from GEO under the accessions cited above. Figure-generation and supplementary-table scripts are retained in the accompanying submission package.")

add_heading(doc, "Figure legends")
for text in (
    "Fig. 1. Multiscale computational framework. Public bulk cohorts support discovery and workflow comparison; locked cohorts, single-cell and spatial data provide independent layers of validation; gutMGene and ancillary docking contextualize lariciresinol prioritization.",
    "Fig. 2. Bulk discovery and candidate convergence. Differential expression, enrichment, WGCNA, glycolysis intersection and cross-cohort PPARG effects define the constrained panel.",
    "Fig. 3. Leakage-controlled workflow comparison and locked external testing of the eight-gene ridge signature.",
    "Fig. 4. Donor-aware single-cell atlas localizes PPARG-associated dysregulation to epithelial cells.",
    "Fig. 5. Spatial transcriptomics provides directional, section-limited support for lower epithelial PPARG regulon activity in UC.",
    "Fig. 6. Immune-cell context and gutMGene evidence support lariciresinol prioritization; database associations are not patient-level measurements.",
):
    add_para(doc, text, size=9.5, after=2)

add_heading(doc, "References")
source_doc = Document(SOURCE)
ref_start = next(i for i, p in enumerate(source_doc.paragraphs) if p.text.strip() == "References")
for p in source_doc.paragraphs[ref_start + 1:]:
    text = p.text.strip()
    if text:
        add_para(doc, text, size=8.8, after=2)

for section in doc.sections:
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        set_font(run, size=run.font.size.pt if run.font.size else 11, bold=run.bold, italic=run.italic)

doc.save(OUTPUT)

body_text = []
for p in doc.paragraphs:
    if p.text.strip() in {"References", "Figure legends"}:
        break
    body_text.append(p.text)
word_count = len(re.findall(r"\b[\w-]+\b", " ".join(body_text)))
print(f"OUTPUT={OUTPUT}")
print(f"MAIN_TEXT_WORDS_EXCL_FIGURE_LEGENDS_AND_REFERENCES={word_count}")
