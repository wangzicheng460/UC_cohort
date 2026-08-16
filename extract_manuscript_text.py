from pathlib import Path
from docx import Document

source = Path(r"C:\Users\wangz\Desktop\初稿_加入Introduction_gutMGeneV2.docx")
doc = Document(source)

for index, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip()
    if text:
        print(f"P{index:04d}\t{paragraph.style.name}\t{text}")

for table_index, table in enumerate(doc.tables):
    print(f"TABLE {table_index}")
    for row_index, row in enumerate(table.rows):
        print(f"R{row_index}\t" + "\t".join(cell.text.replace("\n", " | ") for cell in row.cells))
