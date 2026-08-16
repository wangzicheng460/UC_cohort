from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

SOURCE = Path(r"C:\Users\wangz\Desktop\初稿_加入Introduction_gutMGeneV2.docx")


def font_info(run):
    rpr = run._element.rPr
    fonts = None if rpr is None else rpr.rFonts
    return {
        "name": run.font.name,
        "ascii": None if fonts is None else fonts.get(qn("w:ascii")),
        "hAnsi": None if fonts is None else fonts.get(qn("w:hAnsi")),
        "eastAsia": None if fonts is None else fonts.get(qn("w:eastAsia")),
        "size": None if run.font.size is None else run.font.size.pt,
        "bold": run.bold,
        "italic": run.italic,
    }


doc = Document(SOURCE)
print(f"SOURCE={SOURCE}")
print(f"PARAGRAPHS={len(doc.paragraphs)} TABLES={len(doc.tables)} SECTIONS={len(doc.sections)}")
print("\n=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text.replace("\t", "<TAB>").replace("\n", "<NL>")
    if text.strip() or p.style.name != "Normal":
        print(f"P{i:04d} [{p.style.name}] {text}")
        if text.strip():
            seen = []
            for run in p.runs:
                info = font_info(run)
                key = tuple(info.items())
                if key not in seen:
                    seen.append(key)
                    print(f"    RUNFMT {info}")

print("\n=== TABLES ===")
for ti, table in enumerate(doc.tables):
    print(f"TABLE {ti} rows={len(table.rows)} cols={len(table.columns)} style={table.style.name if table.style else None}")
    for ri, row in enumerate(table.rows):
        vals = [cell.text.replace("\n", " | ") for cell in row.cells]
        print(f"  R{ri}: {vals}")

print("\n=== STYLES ===")
for name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
    try:
        style = doc.styles[name]
    except KeyError:
        continue
    print(name, style.font.name, None if style.font.size is None else style.font.size.pt)
