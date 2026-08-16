import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


path = Path(r"C:\Users\wangz\Desktop\初稿_加入Introduction_gutMGeneV2_完成版.docx")
doc = Document(path)
paragraphs = doc.paragraphs
texts = [p.text.strip() for p in paragraphs]

expected_title = (
    "Integrated transcriptomics identifies a reproducible epithelial PPARG-associated axis "
    "and prioritizes lariciresinol in ulcerative colitis"
)
assert texts[0] == expected_title
for heading in ("4. Discussion", "5. Conclusion", "6. References"):
    assert texts.count(heading) == 1, heading

ref_start = texts.index("6. References") + 1
reference_numbers = []
for text in texts[ref_start:]:
    match = re.match(r"^(\d+)\.\s", text)
    if match:
        reference_numbers.append(int(match.group(1)))
assert reference_numbers == list(range(1, 31)), reference_numbers

cited = set()
for text in texts[: ref_start - 1]:
    for bracket in re.findall(r"\[([0-9,\-\s]+)\]", text):
        for item in bracket.split(","):
            item = item.strip()
            if "-" in item:
                start, end = map(int, item.split("-", 1))
                cited.update(range(start, end + 1))
            elif item:
                cited.add(int(item))
assert cited and min(cited) >= 1 and max(cited) <= 30, sorted(cited)

joined = "\n".join(texts)
for token in ("TODO", "Evidence needed", "4.Discussion", "5.conclusion", "{{", "}}"):
    assert token not in joined, token

runs = []
for paragraph in paragraphs:
    runs.extend(run for run in paragraph.runs if run.text.strip())
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                runs.extend(run for run in paragraph.runs if run.text.strip())
for section in doc.sections:
    for container in (section.header, section.footer):
        for paragraph in container.paragraphs:
            runs.extend(run for run in paragraph.runs if run.text.strip())

bad_fonts = []
for run in runs:
    rfonts = run._element.get_or_add_rPr().rFonts
    names = {run.font.name}
    if rfonts is not None:
        names.update(rfonts.get(qn(key)) for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"))
    names.discard(None)
    if names != {"Times New Roman"}:
        bad_fonts.append((run.text[:40], sorted(names)))
assert not bad_fonts, bad_fonts[:10]

captions = [p for p in paragraphs if p.text.strip().startswith("Fig.")]
assert len(captions) == 5
assert all(p._p.pPr.find(qn("w:wordWrap")) is None for p in captions)
assert len(doc.inline_shapes) == 5

print(f"PASS: {len(paragraphs)} paragraphs; {len(captions)} figures; 30 sequential references")
print(f"PASS: citation range {min(cited)}-{max(cited)}; {len(runs)} non-empty runs all Times New Roman")
