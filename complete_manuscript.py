import re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

SOURCE = Path(r"C:\Users\wangz\Desktop\初稿_加入Introduction_gutMGeneV2.docx")
OUTPUT = Path(r"C:\Users\wangz\Desktop\初稿_加入Introduction_gutMGeneV2_完成版.docx")
FONT = "Times New Roman"


def set_run_font(run, size=None, bold=None, italic=None):
    run.font.name = FONT
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), FONT)


def set_paragraph_text(paragraph, text, size=12, bold=False):
    paragraph.text = ""
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def set_labeled_paragraph(paragraph, label, text):
    paragraph.text = ""
    label_run = paragraph.add_run(label)
    set_run_font(label_run, size=12, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=12, bold=False)


def find_paragraph(doc, startswith):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(startswith):
            return paragraph
    raise ValueError(f"Paragraph not found: {startswith}")


def insert_before(reference_paragraph, text, *, size=12, bold=False, align=None, space_before=0, space_after=6):
    paragraph = reference_paragraph.insert_paragraph_before()
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


doc = Document(SOURCE)

# Add a defensible, finding-led title before the existing keyword line.
title = doc.paragraphs[0].insert_paragraph_before()
title_run = title.add_run(
    "Integrated transcriptomics identifies a reproducible epithelial PPARG-associated axis "
    "and prioritizes lariciresinol in ulcerative colitis"
)
set_run_font(title_run, size=16, bold=True)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(12)
title.paragraph_format.keep_with_next = True

# Correct the discovery-cohort description: GSE107499 uses non-inflamed UC tissue as reference.
abstract_methods = find_paragraph(doc, "Methods:We integrated")
set_labeled_paragraph(
    abstract_methods,
    "Methods: ",
    "We integrated four discovery mucosal transcriptome cohorts (GSE73661, GSE75214, "
    "GSE87466 and GSE107499; 326 UC or inflamed-UC samples and 88 reference samples, "
    "comprising 44 non-IBD controls and 44 non-inflamed UC biopsies), two external validation "
    "cohorts, donor-aware single-cell RNA sequencing, spatial transcriptomics and gutMGene v2.0 "
    "curation. Differential expression analysis, weighted gene co-expression network analysis, "
    "glycolysis-related gene-set integration and 117 feature-selection/classifier workflows with "
    "leave-one-cohort-out validation were used to identify stable candidates. Single-cell and "
    "spatial analyses were then performed to characterize PPARG-associated cellular and metabolic "
    "patterns and to prioritize a gut microbe-linked metabolite."
)

# Normalize spacing after the remaining abstract labels without changing their wording.
for label, prefix in (
    ("Background: ", "Background:Ulcerative"),
    ("Results: ", "Results:The integrative"),
    ("Conclusion: ", "Conclusion:These findings"),
):
    paragraph = find_paragraph(doc, prefix)
    set_labeled_paragraph(paragraph, label, paragraph.text.split(":", 1)[1])

replacements = {
    "We reanalysed publicly available": (
        "We reanalysed publicly available, de-identified transcriptomic datasets from the NCBI "
        "Gene Expression Omnibus (GEO) [9]. Four bulk mucosal cohorts were designated as the discovery "
        "set: GSE73661 (12 non-IBD controls and 67 UC samples) [10], GSE75214 (11 controls and 97 UC "
        "samples) [11], GSE87466 (21 normal controls and 87 UC samples) [12], and GSE107499 (44 "
        "non-inflamed and 75 inflamed biopsies, all from individuals with UC) [13], giving 414 samples "
        "in total. Thus, the aggregated reference group comprised 44 non-IBD controls and 44 "
        "non-inflamed UC biopsies; GSE107499 was not a healthy-control cohort. GSE47908 (15 controls "
        "and 39 UC samples after exclusion of six UC dysplasia samples) [14] and GSE13367 (10 controls "
        "and 17 UC samples after exclusion of collagenous-colitis, isolated-colonocyte and other "
        "non-target samples) [15] were kept locked until final external evaluation."
    ),
    "For cellular validation": (
        "For cellular validation, GSE214695 [16] was restricted to six healthy-control and six UC "
        "donors; Crohn's disease samples were excluded. The retained object contained 30,068 cells "
        "and was analysed with donor-level biological replication. For spatial validation, GSE189184 "
        "[17] contributed seven Visium sections: B10 and C5 (healthy controls) and B4, B5, B12, B13 "
        "and C2 (UC). Section/patient, rather than Visium spot, was treated as the biological "
        "replicate. All analyses used the accession-level sample manifests and exclusions recorded "
        "in the project folder."
    ),
    "Processed expression matrices": (
        "Processed expression matrices were restricted to genes shared across the four discovery "
        "cohorts. Genes detected above zero in at least 10% of samples were retained, values were "
        "transformed as log2(x + 1), and empirical-Bayes ComBat [18] was applied with cohort as the "
        "batch variable for discovery-level integration. Principal-component analyses before and "
        "after correction were used as diagnostics. The corrected matrix contained 16,251 genes."
    ),
    "Differential expression between": (
        "Differential expression between the study-defined target and reference groups was performed "
        "with limma [19] using a no-intercept design and the target-versus-reference contrast. Genes "
        "with Benjamini-Hochberg-adjusted P < 0.05 and absolute log2 fold change > 1 were defined as "
        "DEGs. Gene Ontology (GO) biological-process [20] and Kyoto Encyclopedia of Genes and Genomes "
        "(KEGG) pathway enrichment [21] were calculated from the DEG list and summarized by gene "
        "ratio, gene count and adjusted significance."
    ),
    "WGCNA was applied": (
        "Weighted gene co-expression network analysis (WGCNA) [22] was applied to the discovery "
        "expression matrix with an automatically selected soft-thresholding power. Modules were "
        "identified by dynamic tree cutting with a minimum module size of 60 and deepSplit = 2, and "
        "module eigengenes were related to the study-defined phenotype. Genes were retained as WGCNA "
        "candidates when absolute gene significance exceeded 0.5 and absolute module membership "
        "exceeded 0.8. A 911-gene glycolysis compendium was assembled from the pathway resources used "
        "in the project, including MSigDB, KEGG and Reactome [21,23,24]. The 24-gene panel was defined "
        "by intersecting the DEG, WGCNA and glycolysis sets. A separate 117-gene gutMGene v2.0-related "
        "set [7] was used only to identify the shared PPARG gene shown in Fig. 2E."
    ),
    "PPARG expression was summarized": (
        "PPARG expression was summarized at broad- and fine-cell-type levels, with PPARG-positive "
        "cells defined from detected expression. Glycolysis activity was scored with the MSigDB "
        "HALLMARK_GLYCOLYSIS gene set [23] using UCell [25]; KEGG and Reactome glycolysis sets "
        "[21,24] were retained as sensitivity analyses. Fig. 4D used the median donor-level "
        "PPARG-positive fraction and the median donor-level Hallmark glycolysis score for each broad "
        "cell type in UC, thereby avoiding cell-level pseudo-replication."
    ),
    "Donor-by-cell-type pseudobulk": (
        "Donor-by-cell-type pseudobulk counts were used for expression-level comparisons. PPARG "
        "transcription-factor activity was inferred from DoRothEA A-C confidence regulons [26] with "
        "weighted activity scoring implemented through decoupleR [27]. Activity differences were "
        "evaluated at the donor level and adjusted across the transcription factors tested within each "
        "compartment. Regulon activity was interpreted as a predicted transcriptional footprint, not "
        "direct DNA binding or demonstrated causal regulation."
    ),
    "CIBERSORT fractions were obtained": (
        "CIBERSORT-derived cell fractions [28] were obtained from the project output and linked to the "
        "phenotype manifest for the four discovery cohorts. Because fractions were bounded and "
        "zero-inflated, values were transformed with the arcsine-square-root transformation before "
        "standardized mean differences were calculated. Hedges' g was computed within each cohort "
        "for the study-defined target versus reference group, pooled with a random-effects "
        "meta-analysis and adjusted across cell types with the Benjamini-Hochberg procedure."
    ),
    "xCell scores were calculated": (
        "xCell scores [29] were calculated from the normalized bulk expression matrix. Within each "
        "discovery cohort, PPARG expression and epithelial- or plasma-cell xCell scores were rank "
        "transformed and residualized on study group to obtain group-adjusted partial Spearman "
        "correlations. Cohort estimates were combined on the Fisher z scale with random-effects "
        "pooling and transformed back to Spearman rho with 95% confidence intervals. These analyses "
        "estimate associations between a gene and computational cell scores; they are not "
        "ligand-receptor, CellChat, LIANA, CellPhoneDB or NicheNet analyses."
    ),
    "Gut microbiota-host-gene": (
        "Gut microbiota-host-gene and microbe-metabolite relations were obtained from the gutMGene "
        "v2.0 database [7] and the locally curated project export. The plotted network contained six "
        "E. faecalis-metabolite relations (agmatine, citrulline, L-leucic acid, lariciresinol, "
        "levodopa and tyramine) and one S. salivarius-butyrate relation. The candidate-versus-gutMGene "
        "v2.0 overlap was used to highlight PPARG, and lariciresinol was selected for subsequent "
        "structure-based evaluation because it was the highlighted E. faecalis-linked metabolite."
    ),
    "Unless stated otherwise": (
        "Unless stated otherwise, tests were two-sided. Multiple comparisons were controlled with the "
        "Benjamini-Hochberg procedure [30]. Donors, bulk samples or Visium sections were used as the "
        "biological units appropriate to each analysis; individual cells and spots were not treated "
        "as independent patients. Effect sizes, confidence intervals and replicate-level "
        "distributions were prioritized over isolated P values. Analyses were implemented in R and "
        "Python using the scripts in the project archive; package versions and session information "
        "should be exported before submission."
    ),
    "We first integrated four discovery": (
        "We first integrated four discovery mucosal transcriptome cohorts (GSE73661, GSE75214, "
        "GSE87466 and GSE107499; 326 UC or inflamed-UC samples and 88 reference samples) [10-13] to "
        "identify reproducible disease- or inflammation-associated signals. The reference samples "
        "comprised 44 non-IBD controls and 44 non-inflamed UC biopsies from GSE107499. Differential "
        "expression analysis identified 681 genes at an adjusted P value < 0.05 and an absolute log2 "
        "fold change > 1, including 437 genes upregulated and 244 genes downregulated in the target "
        "group (Fig. 2A). The upregulated markers included LCN2 (log2 fold change 2.680; adjusted "
        "P = 4.17 x 10-67) and SLC2A3 (log2 fold change 1.518; adjusted P = 1.12 x 10-36), whereas "
        "PPARG was strongly reduced (log2 fold change -1.294; adjusted P = 8.26 x 10-42)."
    ),
    "The PPARG decrease was also": (
        "The PPARG decrease was also consistent at the cohort level. Hedges' g for the study-defined "
        "target versus reference contrast was negative in all four discovery cohorts: -2.465 in "
        "GSE73661, -2.086 in GSE75214, -1.974 in GSE87466 and -2.514 in GSE107499; each 95% confidence "
        "interval excluded zero (Fig. 2F). Thus, the direction of PPARG suppression was not driven by "
        "a single discovery study, although the GSE107499 contrast represents inflamed versus "
        "non-inflamed UC rather than UC versus healthy control [13]."
    ),
    "To connect the differential signal": (
        "To connect the differential signal to co-expression structure and metabolism, we related "
        "WGCNA module eigengenes [22] to the study-defined phenotype and selected genes with strong "
        "module membership and gene-significance values. The plotted WGCNA set contained 294 genes, "
        "and the glycolysis compendium contained 911 genes (Fig. 2D,E). Intersecting these sets with "
        "the 681 DEGs produced 24 candidate genes: DDC, TSPAN7, LCN2, RHOU, SOCS3, ANXA1, SERPINA1, "
        "CHST2, FGR, PDE6A, TRPM6, CATSPERB, MOGAT2, SERPING1, KDELR3, ME1, VLDLR, PCK1, MAOA, ADH1C, "
        "ADH6, PFKFB3, SLC2A3 and PPARG."
    ),
    "The GSE214695 single-cell": (
        "The GSE214695 single-cell analysis [16] retained 30,068 cells from six healthy-control and "
        "six UC donors after excluding Crohn's disease samples and collapsed the author annotations "
        "into ten broad compartments (Fig. 4A). Donor-level proportions showed a marked loss of "
        "epithelial cells in UC, with a mean UC-HC difference of -37.47 percentage points (bootstrap "
        "95% CI -49.53 to -24.44; FDR = 3.97 x 10-4) (Fig. 4B). In parallel, plasma cells increased "
        "by 24.74 percentage points (FDR = 0.00935), myeloid cells by 3.25 percentage points "
        "(FDR = 0.0120) and neutrophils by 0.95 percentage points (FDR = 0.00754). Enteric glia also "
        "decreased by 0.36 percentage points (FDR = 5.96 x 10-4). Donors, rather than individual "
        "cells, were treated as the biological replicates."
    ),
    "We next evaluated seven Visium": (
        "We next evaluated seven Visium sections from GSE189184 [17]: two healthy-control sections "
        "(B10 and C5) and five active-UC sections (B4, B5, B12, B13 and C2). After spot-level quality "
        "control, 8,306 spots were retained. Deconvolution of the spatial mixtures reproduced the "
        "direction of the single-cell composition changes, including lower epithelial signal and "
        "higher plasma-cell, B-cell, neutrophil and myeloid signal in UC (Fig. 5A,B). Exact two-sided "
        "section-level tests were not significant, reflecting the limited number of control sections."
    ),
    "Across the four bulk discovery": (
        "Across the four bulk discovery cohorts, CIBERSORT-derived immune-cell estimates [28] showed "
        "broad but heterogeneous alterations in the study-defined target group (Fig. 6A). Random-effects "
        "pooled Hedges' g values were positive for macrophages M0 (+2.117; FDR = 0.00141), activated "
        "CD4 memory T cells (+1.079; FDR < 0.001), neutrophils (+0.890; FDR < 0.001), activated mast "
        "cells (+0.726; FDR = 0.00321), M1 macrophages (+1.076; FDR = 0.0225) and resting NK cells "
        "(+1.079; FDR = 0.0467). Negative pooled effects were observed for M2 macrophages (-2.208; "
        "FDR < 0.001), resting mast cells (-0.853; FDR = 0.0225), CD8 T cells (-0.619; FDR = 0.0142) "
        "and regulatory T cells (-0.594; FDR = 0.0174). The plasma-cell estimate was positive but "
        "imprecise (Hedges' g = +0.259; FDR = 0.651)."
    ),
    "Disease-adjusted partial Spearman": (
        "Study-group-adjusted partial Spearman analysis of PPARG expression and xCell scores [29] "
        "showed a reproducible positive association with epithelial-cell scores and a negative "
        "association with plasma-cell scores in every discovery cohort (Fig. 6B). The pooled epithelial "
        "association was rho = 0.497 (95% CI 0.413-0.572), whereas the pooled plasma-cell association "
        "was rho = -0.473 (95% CI -0.645 to -0.255). These results are statistical associations between "
        "PPARG expression and computationally estimated cell scores; they do not demonstrate that "
        "PPARG regulates infiltration or that plasma cells communicate directly with epithelial cells."
    ),
    "The gutMGene v2.0-derived": (
        "The gutMGene v2.0-derived network [7] supplied a traceable microbe-metabolite context for the "
        "PPARG hypothesis (Fig. 6C). The plotted relations linked Enterococcus faecalis to agmatine, "
        "citrulline, L-leucic acid, lariciresinol, levodopa and tyramine, and linked Streptococcus "
        "salivarius to butyrate. Lariciresinol was highlighted because it was the selected candidate "
        "for downstream PPARG structure-based testing and because an independent bacterial study "
        "reported transformation of (+)-pinoresinol to (+)-lariciresinol by an E. faecalis strain [8]."
    ),
}

for prefix, replacement in replacements.items():
    paragraph = find_paragraph(doc, prefix)
    set_paragraph_text(paragraph, replacement, size=12, bold=False)

# Update the discovery summary table to reflect the actual GSE107499 reference group.
table = doc.tables[0]
table.cell(1, 1).text = "Discovery bulk; reference: 44 non-IBD controls + 44 non-inflamed UC biopsies; target: 326 UC/inflamed-UC samples"
table.cell(1, 3).text = "DEG, GO/KEGG, WGCNA, candidate panel and LOCO workflow comparison"

# Fill the existing Discussion and Conclusion placeholders, then restore a References heading.
discussion_heading = find_paragraph(doc, "4.Discussion")
conclusion_heading = find_paragraph(doc, "5.conclusion")
first_reference = find_paragraph(doc, "1. Halfvarson")

# Remove empty placeholder paragraphs between the supplied section headings.
paragraphs_now = list(doc.paragraphs)
discussion_index = next(i for i, p in enumerate(paragraphs_now) if p._element is discussion_heading._element)
conclusion_index = next(i for i, p in enumerate(paragraphs_now) if p._element is conclusion_heading._element)
reference_index = next(i for i, p in enumerate(paragraphs_now) if p._element is first_reference._element)
for paragraph in paragraphs_now[discussion_index + 1 : conclusion_index] + paragraphs_now[conclusion_index + 1 : reference_index]:
    if not paragraph.text.strip():
        paragraph._element.getparent().remove(paragraph._element)

set_paragraph_text(discussion_heading, "4. Discussion", size=14, bold=True)
discussion_heading.paragraph_format.keep_with_next = True
discussion_heading.paragraph_format.space_before = Pt(12)
discussion_heading.paragraph_format.space_after = Pt(6)

set_paragraph_text(conclusion_heading, "5. Conclusion", size=14, bold=True)
conclusion_heading.paragraph_format.keep_with_next = True
conclusion_heading.paragraph_format.space_before = Pt(12)
conclusion_heading.paragraph_format.space_after = Pt(6)

references_heading = first_reference.insert_paragraph_before()
set_paragraph_text(references_heading, "6. References", size=14, bold=True)
references_heading.paragraph_format.keep_with_next = True
references_heading.paragraph_format.space_before = Pt(12)
references_heading.paragraph_format.space_after = Pt(6)

discussion_paragraphs = [
    (
        "This study identifies a reproducible PPARG-associated signal in ulcerative-colitis mucosa "
        "and resolves its principal cellular context. Across four discovery cohorts, PPARG was "
        "consistently lower in the study-defined disease or inflamed group, ranked second by feature-"
        "selection stability, and contributed to an eight-gene model that retained strong discrimination "
        "in two locked external cohorts. Donor-aware single-cell analysis localized PPARG mainly to the "
        "epithelial compartment and showed significantly reduced predicted epithelial PPARG regulon "
        "activity in UC. Spatial analysis reproduced the same direction, although section-level tests "
        "were not significant. Together, these results support a suppressed epithelial PPARG-associated "
        "transcriptional axis in UC, while stopping short of demonstrating direct control of glycolysis "
        "or microbial-metabolite regulation of PPARG."
    ),
    (
        "The cross-cohort decrease in PPARG extends earlier evidence that epithelial PPARG expression "
        "is impaired in UC [5] and is consistent with experimental studies in which PPARG agonism "
        "attenuated epithelial inflammatory responses and colitis severity [4]. The present analysis "
        "adds two forms of resolution: reproducibility across heterogeneous bulk cohorts and localization "
        "to an epithelial transcriptional footprint in donor-level single-cell data. The stronger case-"
        "control separation of regulon activity than of epithelial PPARG expression may indicate that "
        "coordinated target-gene changes capture pathway disruption more sensitively than a single, "
        "sparsely detected transcript. However, the activity score depends on a curated DoRothEA network "
        "and a computational footprint model [26,27]; it therefore cannot be interpreted as direct PPARG "
        "binding or causal transcriptional regulation."
    ),
    (
        "The machine-learning analysis also illustrates why cohort structure matters in transcriptomic "
        "biomarker studies. Comparing 117 workflows within outer leave-one-cohort-out folds limited "
        "leakage during feature selection, scaling and tuning, and the locked eight-gene ridge model "
        "outperformed the 24-gene model and PPARG alone in both external cohorts. Nevertheless, the "
        "discovery panel itself was defined using all discovery cohorts, so the leave-one-cohort-out AUCs "
        "are best viewed as workflow-comparison estimates rather than fully independent validation. The "
        "external cohorts were small, particularly GSE13367 (n = 27), and their confidence intervals and "
        "calibration estimates were correspondingly broad. The model should therefore be regarded as a "
        "reproducible molecular signature, not a clinically deployable diagnostic test, until it is "
        "evaluated prospectively in larger, treatment- and activity-stratified cohorts."
    ),
    (
        "The cell-resolved results place the PPARG signal within the extensive epithelial and immune "
        "remodelling previously described in UC [2,6,16]. The marked reduction in epithelial-cell "
        "proportion and increase in plasma cells provide a parsimonious explanation for the positive "
        "PPARG-epithelial xCell association and the inverse PPARG-plasma-cell association observed in "
        "bulk tissue. These associations therefore need not imply epithelial-plasma signalling or PPARG-"
        "dependent immune recruitment. The glycolysis findings are similarly bounded. Prior work linked "
        "PFKFB3-mediated glycolysis to a pro-inflammatory stromal-fibroblast phenotype in IBD [3], whereas "
        "the present study observed PPARG and glycolysis co-enrichment most strongly in epithelial cell "
        "types. Yet neither donor-level PPARG-glycolysis correlations nor broad-cell-type case-control "
        "glycolysis differences survived significance testing. The data support co-localization of two "
        "metabolic features, not mediation or regulatory direction."
    ),
    (
        "Spatial transcriptomics supplied an anatomical layer of validation but also defined an important "
        "boundary. PPARG regulon activity was lower in UC across all three prespecified spatial summaries, "
        "and PPARG activity correlated positively with glycolysis within epithelial-enriched spots. However, "
        "only two control and five UC sections were available, making the exact two-sided P value no smaller "
        "than 0.143 for the observed group allocation. In addition, spot-level correlations are descriptive "
        "because neighbouring spots are spatially autocorrelated. The spatial results are thus directionally "
        "consistent with the single-cell findings but do not provide independent statistical or mechanistic "
        "confirmation."
    ),
    (
        "The gutMGene v2.0 analysis should be interpreted as candidate prioritization. The database provides "
        "a curated E. faecalis-lariciresinol relation [7], and an independent biochemical study showed that "
        "an E. faecalis strain can transform (+)-pinoresinol into (+)-lariciresinol [8]. Neither source shows "
        "that E. faecalis abundance, lariciresinol exposure or PPARG binding differs in the patients analysed "
        "here. Establishing the proposed link will require matched microbiome and metabolomics measurements, "
        "biochemical binding or reporter assays, and PPARG perturbation in epithelial systems. Docking and "
        "molecular dynamics may refine a structural hypothesis, but experimental target engagement and "
        "functional rescue will be necessary before therapeutic relevance can be claimed."
    ),
    (
        "Several limitations constrain generalization. First, this is a retrospective reanalysis of public "
        "datasets that differ in platform, treatment exposure, disease activity and sampling design. In "
        "particular, the 44 reference biopsies in GSE107499 are non-inflamed UC tissue rather than healthy "
        "controls [13], so that cohort contributes evidence about inflammation within UC, not an independent "
        "UC-versus-healthy contrast. Second, residual batch and composition effects may remain despite "
        "normalization, and the external cohorts are too small for definitive calibration. Third, the "
        "single-cell analysis included 12 donors and the spatial analysis seven sections, limiting power for "
        "donor-level correlations and group comparisons. Fourth, deconvolution, gene-set scores and regulon "
        "activities are model-based summaries rather than direct measurements. Finally, the cohorts lack "
        "matched microbial, metabolite and protein-level data. Replication in larger cohorts with matched "
        "mucosal transcriptomics, microbiome and metabolomics, followed by epithelial PPARG perturbation, "
        "would provide the clearest test of the proposed host-microbe-metabolic axis."
    ),
]

for text in discussion_paragraphs:
    insert_before(conclusion_heading, text, size=12, space_after=6)

insert_before(
    references_heading,
    "By integrating multi-cohort bulk transcriptomics with leakage-controlled modelling, donor-aware "
    "single-cell analysis and spatial validation, this study identifies a reproducibly suppressed, "
    "predominantly epithelial PPARG-associated axis in ulcerative colitis. The decisive evidence is the "
    "consistent cross-cohort PPARG decrease together with significantly reduced predicted epithelial PPARG "
    "regulon activity; the spatial and glycolysis analyses provide directional or co-localization evidence "
    "rather than proof of mechanism. gutMGene v2.0 and prior bacterial biotransformation data further "
    "prioritize lariciresinol as a tractable hypothesis, but direct PPARG binding, microbiome-mediated "
    "regulation and therapeutic effects remain untested. These boundaries position the PPARG-lariciresinol "
    "link as a focused target for matched multi-omics and experimental validation rather than as an "
    "established causal pathway.",
    size=12,
    space_after=10,
)

# Append verified references cited in Methods, Results, Discussion and Conclusion.
references = [
    "9. Clough E, Barrett T, Wilhite SE, et al. NCBI GEO: archive for gene expression and epigenomics data sets: 23-year update. Nucleic Acids Research. 2024;52(D1):D138-D144. doi:10.1093/nar/gkad965.",
    "10. Arijs I, De Hertogh G, Lemmens B, et al. Effect of vedolizumab (anti-alpha4beta7-integrin) therapy on histological healing and mucosal gene expression in patients with ulcerative colitis. Gut. 2018;67(1):43-52. doi:10.1136/gutjnl-2016-312293.",
    "11. Vancamelbeke M, Vanuytsel T, Farre R, et al. Genetic and transcriptomic bases of intestinal epithelial barrier dysfunction in inflammatory bowel disease. Inflammatory Bowel Diseases. 2017;23(10):1718-1729. doi:10.1097/MIB.0000000000001246.",
    "12. Li K, Strauss R, Ouahed J, et al. Molecular comparison of adult and pediatric ulcerative colitis indicates broad similarity of molecular pathways in disease tissue. Journal of Pediatric Gastroenterology and Nutrition. 2018;67(1):45-52. doi:10.1097/MPG.0000000000001898.",
    "13. National Center for Biotechnology Information. GSE107499: Expression data from Ulcerative Colitis subjects. Gene Expression Omnibus. https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE107499. Accessed August 16, 2026.",
    "14. Bjerrum JT, Nielsen OH, Riis LB, et al. Transcriptional analysis of left-sided colitis, pancolitis, and ulcerative colitis-associated dysplasia. Inflammatory Bowel Diseases. 2014;20(12):2340-2352. doi:10.1097/MIB.0000000000000235.",
    "15. Bjerrum JT, Hansen M, Olsen J, Nielsen OH. Genome-wide gene expression analysis of mucosal colonic biopsies and isolated colonocytes suggests a continuous inflammatory state in the lamina propria of patients with quiescent ulcerative colitis. Inflammatory Bowel Diseases. 2010;16(6):999-1007. doi:10.1002/ibd.21142.",
    "16. Garrido-Trigo A, Corraliza AM, Veny M, et al. Macrophage and neutrophil heterogeneity at single-cell spatial resolution in human inflammatory bowel disease. Nature Communications. 2023;14:4506. doi:10.1038/s41467-023-40156-6.",
    "17. Gupta T, Antanaviciute A, Lee CHJ, et al. Tracking in situ checkpoint inhibitor-bound target T cells in patients with checkpoint-induced colitis. Cancer Cell. 2024;42(5):797-814.e15. doi:10.1016/j.ccell.2024.04.010.",
    "18. Johnson WE, Li C, Rabinovic A. Adjusting batch effects in microarray expression data using empirical Bayes methods. Biostatistics. 2007;8(1):118-127. doi:10.1093/biostatistics/kxj037.",
    "19. Ritchie ME, Phipson B, Wu D, et al. limma powers differential expression analyses for RNA-sequencing and microarray studies. Nucleic Acids Research. 2015;43(7):e47. doi:10.1093/nar/gkv007.",
    "20. Gene Ontology Consortium. The Gene Ontology resource: enriching a GOld mine. Nucleic Acids Research. 2021;49(D1):D325-D334. doi:10.1093/nar/gkaa1113.",
    "21. Kanehisa M, Goto S. KEGG: Kyoto Encyclopedia of Genes and Genomes. Nucleic Acids Research. 2000;28(1):27-30. doi:10.1093/nar/28.1.27.",
    "22. Langfelder P, Horvath S. WGCNA: an R package for weighted correlation network analysis. BMC Bioinformatics. 2008;9:559. doi:10.1186/1471-2105-9-559.",
    "23. Liberzon A, Birger C, Thorvaldsdottir H, Ghandi M, Mesirov JP, Tamayo P. The Molecular Signatures Database hallmark gene set collection. Cell Systems. 2015;1(6):417-425. doi:10.1016/j.cels.2015.12.004.",
    "24. Milacic M, Beavers D, Conley P, et al. The Reactome Pathway Knowledgebase 2024. Nucleic Acids Research. 2024;52(D1):D672-D678. doi:10.1093/nar/gkad1025.",
    "25. Andreatta M, Carmona SJ. UCell: robust and scalable single-cell gene signature scoring. Computational and Structural Biotechnology Journal. 2021;19:3796-3798. doi:10.1016/j.csbj.2021.06.043.",
    "26. Garcia-Alonso L, Holland CH, Ibrahim MM, Turei D, Saez-Rodriguez J. Benchmark and integration of resources for the estimation of human transcription factor activities. Genome Research. 2019;29(8):1363-1375. doi:10.1101/gr.240663.118.",
    "27. Badia-i-Mompel P, Velez Santiago J, Braunger J, et al. decoupleR: ensemble of computational methods to infer biological activities from omics data. Bioinformatics Advances. 2022;2(1):vbac016. doi:10.1093/bioadv/vbac016.",
    "28. Newman AM, Liu CL, Green MR, et al. Robust enumeration of cell subsets from tissue expression profiles. Nature Methods. 2015;12(5):453-457. doi:10.1038/nmeth.3337.",
    "29. Aran D, Hu Z, Butte AJ. xCell: digitally portraying the tissue cellular heterogeneity landscape. Genome Biology. 2017;18:220. doi:10.1186/s13059-017-1349-1.",
    "30. Benjamini Y, Hochberg Y. Controlling the false discovery rate: a practical and powerful approach to multiple testing. Journal of the Royal Statistical Society Series B. 1995;57(1):289-300. doi:10.1111/j.2517-6161.1995.tb02031.x.",
]

for reference in references:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(reference)
    set_run_font(run, size=11)

# Apply Times New Roman throughout, including tables, headers and footers.
for style in doc.styles:
    if hasattr(style, "font"):
        style.font.name = FONT
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:cs"), FONT)

doc.styles["Normal"].font.size = Pt(12)

for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        set_run_font(run)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=10)

for section in doc.sections:
    for container in (section.header, section.footer):
        for paragraph in container.paragraphs:
            for run in paragraph.runs:
                set_run_font(run)

# Prevent Word's distributed caption alignment from splitting English words across lines.
for paragraph in doc.paragraphs:
    if paragraph.text.strip().startswith("Fig."):
        paragraph.style = doc.styles["Normal"]
        word_wrap = paragraph._p.pPr.find(qn("w:wordWrap"))
        if word_wrap is not None:
            paragraph._p.pPr.remove(word_wrap)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.keep_with_next = False
        for run in paragraph.runs:
            set_run_font(run, size=10.5)

# Use one compact, readable format for the complete numbered reference list.
in_references = False
for paragraph in doc.paragraphs:
    if paragraph.text.strip() == "6. References":
        in_references = True
        continue
    if in_references and re.match(r"^\d+\.\s", paragraph.text.strip()):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_after = Pt(3)
        for run in paragraph.runs:
            set_run_font(run, size=11)

doc.save(OUTPUT)
print(OUTPUT)
